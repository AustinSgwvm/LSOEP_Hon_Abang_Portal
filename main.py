# ==============================================================================
# 🏛️ LSOEP PORTAL PLATFORM ENGINE - INTEGRATED MASTER ROUTER
# Project: Ikom/Boki Federal Constituency (Honourable Victor Abang, PhD)
# File: main.py (V75.1 - Cloud-Compatible Runtime Path Patch)
# ==============================================================================

import sys
import os
import asyncio
import warnings
import streamlit as st
import base64
from docx import Document

# --- 1. SUPER-EARLY STATE INITIALIZATION ---
if "current_route" not in st.session_state:
    st.session_state.current_route = "HOME"
if "admin_authenticated" not in st.session_state:
    st.session_state.admin_authenticated = False
if "agent_authenticated" not in st.session_state:
    st.session_state.agent_authenticated = False
if "admin_module_view" not in st.session_state:
    st.session_state.admin_module_view = "📊 Master Registry Matrix"
if "agent_active_tier" not in st.session_state:
    st.session_state.agent_active_tier = "🦅 Presidential"

# --- IMMUTABLE BASE STATE SEEDING FOR NATIONWIDE MATRIX ---
if "DYNAMIC_GEO_MATRIX" not in st.session_state:
    st.session_state.DYNAMIC_GEO_MATRIX = {
        "Gombe": {
            "Akko": ["Kumo Central", "Kumo East", "Kumo West", "Pindiga", "Gona", "Kaltungo"],
            "Balanga": ["Bambam", "Bangu", "Dadiya", "Galam", "Tal", "Siri", "Mwona"],
            "Billiri": ["Billiri-North", "Billiri-South", "Bare", "Kantali", "Tanglang", "Todi"],
            "Dukku": ["Dukku", "Gombe Abba", "Malala", "Zaune", "Hashidu"],
        },
        "FCT": {
            "AMAC": ["Garki", "Wuse", "Asokoro", "Maitama", "Gwarinpa", "Karshi", "Kabusa"],
            "Gwagwalada": ["Central", "Staff Quarters", "Ikwa", "Paiko", "Dobi", "Zuba"],
        },
        "Cross River": {
            "Ikom": ["Ikom Urban", "Olulumo", "Ofutop I", "Ofutop II", "Nta/Selimba", "Abanyom", "Yala"],
            "Boki": ["Boki East", "Boki West", "Boki North", "Boki South", "Osokom", "Wula", "Boje"],
            "Ogoja": ["Ogoja Urban", "Mbube I", "Mbube II", "Ekajuk", "Nkum I", "Nkum II"],
            "Calabar Municipal": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5", "Ward 6", "Ward 7"],
        },
        "Abia": {
            "Aba North": ["Ward 1", "Ward 2", "Ward 3", "Eziama", "Industrial"],
            "Aba South": ["Ward 4", "Ward 5", "Aba River", "Asa", "Enyimba"],
            "Arochukwu": ["Arochukwu I", "Arochukwu II", "Ohafia", "Ututu"],
            "Umuahia North": ["Urban I", "Urban II", "Afugiri", "Ibeku East", "Isieke"],
            "Umuahia South": ["Amachara", "Olokoro I", "Olokoro II", "Ubakala A", "Ubakala B"]
        },
        "Adamawa": {
            "Yola North": ["Alkalawa", "Doueli", "Gwadabawa", "Nassarawo", "Ajiya"],
            "Yola South": ["Adarawo", "Bole", "Bako", "Mbamba", "Yoldewali"],
            "Mubi North": ["Kolere", "Sabon Gari", "Yelwa", "Digil", "Lokuwa"],
            "Mubi South": ["Gude", "Nassarawo", "Lamurde", "Mugulbu", "Yele"]
        },
        "Akwa Ibom": {
            "Uyo": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5", "Ward 6"],
            "Eket": ["Urban I", "Urban II", "Urban III", "Central I", "Central II"],
            "Ikot Ekpene": ["Urban I", "Urban II", "Amanya", "Ikot Inyang", "Mbiaso"],
            "Oron": ["Oron Urban I", "Oron Urban II", "Oron Urban III", "Oron Urban IV"]
        },
        "Anambra": {
            "Awka South": ["Ward 1", "Ward 2", "Ward 3", "Amawbia I", "Amawbia II"],
            "Onitsha North": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5"],
            "Onitsha South": ["Ward 6", "Ward 7", "Ward 8", "Fegge I", "Fegge II"],
            "Nnewi North": ["Otolo I", "Otolo II", "Uruagu I", "Uruagu II", "Umudim"]
        },
        "Bauchi": {
            "Bauchi LGA": ["Majema", "Makama", "Dankade", "Hardawa", "Zunjil"],
            "Katagum": ["Azare", "Chinade", "Madara", "Busau", "Gambaki"],
            "Misau": ["Misau Central", "Misau East", "Ajigin", "Gwaram"],
            "Alkaleri": ["Alkaleri", "Pali", "Gwana", "Mansur", "Yali"]
        },
        "Bayelsa": {
            "Yenagoa": ["Epie I", "Epie II", "Gbarain", "Okordia", "Zarama"],
            "Brass": ["Ward 1", "Ward 2", "Ward 3", "Cape Formosa"],
            "Sagbama": ["Sagbama Towns", "Agbere", "Toro-Orua", "Ofoni"],
            "Southern Ijaw": ["Oporoma", "Olah", "Amassoma", "Ekowe"]
        },
        "Benue": {
            "Makurdi": ["Central", "North", "South", "Fiidi", "Wailomayo"],
            "Otukpo": ["Town East", "Town West", "Otukpo Central", "Adoka"],
            "Gboko": ["Gboko Central", "Gboko East", "Gboko Northwest", "Yandev"],
            "Katsina-Ala": ["Katsina-Ala Township", "Manger", "Tiir", "Utange"]
        },
        "Borno": {
            "Maiduguri": ["Shehuri", "Maisandari", "Bolori I", "Bolori II", "Gwange"],
            "Biu": ["Biu Central", "Biu East", "Miringa", "Buratai"],
            "Gwoza": ["Gwoza Towns", "Pulka", "Ashigashiya", "Limankara"],
            "Bama": ["Bama Towns", "Kasugula", "Soyoye"]
        },
        "Delta": {
            "Asaba": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5"],
            "Warri South": ["Urban I", "Urban II", "Urban III", "Odeh-Itsekiri"],
            "Uvwie": ["Effurun I", "Effurun II", "Enerhen", "Ekpan"],
            "Ughelli North": ["Ughelli Urban I", "Ughelli Urban II", "Agbarho", "Evwreni"]
        },
        "Ebonyi": {
            "Abakaliki": ["Azuiyi", "Azugwu", "Urban I", "Urban II", "Ameke"],
            "Afikpo North": ["Oziza", "Amisu", "Unwana I", "Unwana II"],
            "Onicha": ["Isu", "Onicha Towns", "Abaomege", "Ukawu"],
            "Ikwo": ["Echara", "Alike", "Noyo", "Umuogudu"]
        },
        "Edo": {
            "Oredo": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5", "Ward 6"],
            "Ikpoba Okha": ["Ward 1", "Ward 2", "Ward 3", "Idogbo", "Aduwawa"],
            "Egor": ["Uselu I", "Uselu II", "Evbutubu", "Ogbeson"],
            "Esan West": ["Ekpoma Towns", "Ihuelen", "Urohi", "Uhiele"]
        },
        "Ekiti": {
            "Ado Ekiti": ["Ado I", "Ado II", "Ado III", "Ado IV", "Ado V"],
            "Ikole": ["Ikole West", "Ikole East", "Odo Oro", "Asin"],
            "Oye": ["Oye Towns", "Iluse", "Ayede", "Ijelu"],
            "Irepodun/Ifelodun": ["Iyin I", "Iyin II", "Igede I", "Igede II"]
        },
        "Enugu": {
            "Enugu North": ["Asata", "Ogui", "China Town", "New Haven"],
            "Enugu South": ["Uwani", "Achara", "Amechi", "Awkunanaw"],
            "Nsukka": ["Nsukka Towns", "Alor-Uno", "Eha-Alumona", "Obukpa"],
            "Oji River": ["Oji Towns", "Inyi", "Achi", "Ugwuoba"]
        },
        "Imo": {
            "Owerri Municipal": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5"],
            "Orlu": ["Central", "East", "Orlu Towns", "Amaifeke", "Owerri-Ebeiri"],
            "Okigwe": ["Okigwe Towns", "Ihube", "Amuro", "Osu"],
            "Mbaitoli": ["Ogwa", "Ubomiri", "Ifakala", "Mbieri"]
        },
        "Jigawa": {
            "Dutse": ["Dutse Takur", "Limawa", "Kudai", "Chamo"],
            "Hadejia": ["Matsaro", "Sabon Garu", "Hadejia Central", "Kafin Hausa"],
            "Kazaure": ["Kazaure Towns", "Dabi", "Roni", "Yankwashi"],
            "Gumel": ["Gumel Towns", "Gagaran", "Hammado", "Zango"]
        },
        "Kaduna": {
            "Kaduna North": ["Shaba", "Gaji", "Kawo", "Dadi", "Unguwan Rimi"],
            "Kaduna South": ["Tudun Wada", "Unguwan Sanusi", "Sabon Gari", "Makera"],
            "Zaria": ["Zaria City", "Tudun Wada", "Gyallesu", "Samaru"],
            "Sabon Gari": ["Sabon Gari Central", "Hanwa", "Dogarawa", "Samaru South"]
        },
        "Kano": {
            "Fagge": ["Fagge North", "Fagge South", "Kwaciri", "Sabon Gari"],
            "Dala": ["Dala Central", "Dogon Nama", "Gwangwazo", "Yakup"],
            "Nassarawa": ["Nassarawa Towns", "Bompai", "Gwagwarwa", "Tudun Wada"],
            "Tarauni": ["Tarauni Towns", "Hotoro", "Unguwa Uku", "Gyadi-Gyadi"]
        },
        "Katsina": {
            "Katsina LGA": ["Wakilin Central", "Wakilin South", "Wakilin North", "Wakilin East"],
            "Daura": ["Daura Arena", "Kofar Baru", "Madobi", "Sarkin Yaki"],
            "Funtua": ["Funtua Towns", "Dandume", "Bakori", "Maska"],
            "Malumfashi": ["Malumfashi Towns", "Yamel", "Dankanjiba", "Kuki"]
        },
        "Kebbi": {
            "Birnin Kebbi": ["Nassarawa", "Rafin Atiku", "Gwandu", "Ambursa"],
            "Argungu": ["Kokani North", "Kokani South", "Felande", "Gulma"],
            "Yauri": ["Yauri Towns", "Gungun", "Zamare", "Tungan"],
            "Zuru": ["Zuru Towns", "Dabag", "Manga", "Rikoto"]
        },
        "Kogi": {
            "Lokoja": ["Ward A", "Ward B", "Ward C", "Ward D", "Ward E"],
            "Okene": ["Bariki", "Onyukoko", "Okene Central", "Odeh"],
            "Anyigba": ["Anyigba Central", "Agbeji", "Egume", "Ologba"],
            "Kabba/Bunu": ["Kabba Towns", "Bunu", "Olle", "Asuta"]
        },
        "Kwara": {
            "Ilorin West": ["Ajikobi", "Baboko", "Adewole", "Oloje", "Warah"],
            "Ilorin East": ["Balogun", "Gambari", "Zango", "Oke-Oyi"],
            "Offa": ["Offa Towns", "Balogun", "Shawo", "Essun"],
            "Edu": ["Lafiagi", "Tsaragi", "Tsonga", "Kpada"]
        },
        "Lagos": {
            "Alimosho": ["Ikotun", "Egbeda", "Ipaja", "Ayobo", "IDimu"],
            "Ikeja": ["Anifowoshe", "Gra", "Oregun", "Alausa", "Ipodo"],
            "Lagos Island": ["Olowogbowo", "Isale Eko", "Lafiaji", "Obalende"],
            "Surulere": ["Ojuelegba", "Adeniran Ogunsanya", "Ijesha", "Masha"]
        },
        "Nasarawa": {
            "Lafia": ["Lafia Central", "Lafia East", "Lafia North", "Assakio"],
            "Karu": ["Mararaba", "Karu Towns", "Ado", "Masaka", "New Karu"],
            "Keffi": ["Keffi Towns", "Yelwa", "Angwan", "Sabon Gari"],
            "Akwanga": ["Akwanga Towns", "Gudi", "Moroa", "Ancho"]
        },
        "Niger": {
            "Minna": ["Central", "Sabon Gari", "Chanchaga", "Bosso"],
            "Bida": ["Landzun", "Masaga", "Dokodza", "Cheniyan"],
            "Suleja": ["Suleja Towns", "Maje", "Iku", "Bwari"],
            "Kontagora": ["Kontagora Central", "Post Office", "Nagwamatse", "Rafin Inya"]
        },
        "Ogun": {
            "Abeokuta South": ["Ake I", "Ake II", "Imo", "Sodeke", "Ijemo"],
            "Ijebu Ode": ["Ijebu North", "Ijebu South", "Isiwo", "Obalende"],
            "Ado-Odo/Ota": ["Ota Towns", "Ado-Odo", "Igbesa", "Agbara"],
            "Sagamu": ["Sagamu Central", "Sabon Gari", "Makun", "Offin"]
        },
        "Ondo": {
            "Akure South": ["Gbogi", "Isinkan", "Oja Oshodi", "Arakale"],
            "Ondo West": ["Urban I", "Urban II", "Urban III", "Yaba", "Idishin"],
            "Owo": ["Owo Towns", "Ehin Ogbe", "Igboroko", "Isaipen"],
            "Ikare": ["Ikare Central", "Okorun", "Iku", "Iyame"]
        },
        "Osun": {
            "Osogbo": ["Alekuwodo", "Ataoja", "Oja Oba", "Okefia"],
            "Ife Central": ["Ilare", "More", "Iremo", "Yekemi"],
            "Ilesa East": ["Ilesa Towns", "Omofe", "Biladu", "Ondo Road"],
            "Ede North": ["Ede Towns", "Oloba", "Sabon Gari", "Bara"]
        },
        "Oyo": {
            "Ibadan North": ["Ward 1", "Ward 2", "Ward 3", "Ward 4", "Ward 5"],
            "Ogbomoso North": ["Isale", "Sabon Gari", "Tara", "Okelerin"],
            "Oyo West": ["Oyo Towns", "Isale Oyo", "Akesan", "Ashipa"],
            "Iseyin": ["Iseyin Towns", "Oke-Oja", "Isalu", "Koso"]
        },
        "Plateau": {
            "Jos North": ["Vanderpuye", "Tafawa Balewa", "Garba Daho", "Gangare"],
            "Jos South": ["Bukuru", "Gyandobolo", "Du", "Vwang"],
            "Pankshin": ["Pankshin Towns", "Lardang", "Wokkos", "Dei"],
            "Shendam": ["Shendam Towns", "Moelet", "Kalong", "Garkawa"]
        },
        "Rivers": {
            "Port Harcourt": ["Diobu", "Town", "Borokiri", "D-Line", "Amadi-Ama"],
            "Obio/Akpor": ["Rumuomasi", "Rumuokwuta", "Rumuigbo", "Elelenwo"],
            "Bonny": ["Bonny Towns", "Finima", "Grand Bonny", "Abalamie"],
            "Ahoada East": ["Ahoada Towns", "Ogbo", "Edeoha", "Upata"]
        },
        "Sokoto": {
            "Sokoto North": ["Waziri A", "Waziri B", "Shehuri", "Magajin Gari"],
            "Sokoto South": ["Sarkin Adar", "Rijiyar Dorowa", "Gagi", "Maitasane"],
            "Wamako": ["Wamako Towns", "Gidan", "Arkilla", "Kalambaina"],
            "Tambuwal": ["Tambuwal Towns", "Dogondaji", "Garki", "Jabo"]
        },
        "Taraba": {
            "Jalingo": ["Turaki A", "Turaki B", "Sarki", "Barade"],
            "Wukari": ["Hospital Ward", "Avyi", "Puje", "Chinkai"],
            "Bali": ["Bali Towns", "Suntai", "Mayo", "Garba"]
        },
        "Yobe": {
            "Damaturu": ["Central", "Nayi-Nawa", "Sabon Gari", "Gwange"],
            "Potiskum": ["Bolewa", "Hausawa", "Yerimara", "Mamudo"],
            "Gashua": ["Gashua Towns", "Sabon Gari", "Sarkin Harki"],
            "Bade": ["Bade Towns", "Gorgoram", "Katamma", "Tagali"]
        },
        "Zamfara": {
            "Gusau": ["Central", "Sabon Gari", "Tudun Wada", "Ruwan Doruwa"],
            "Kaura Namoda": ["Bangana", "Sabon Gari", "Kaura Towns", "Damba"],
            "Maradun": ["Maradun Towns", "Gora", "Damaga", "Dosara"],
            "Talata Mafara": ["Mafara Towns", "Garbadu", "Matusgi", "Jangebe"]
        }
    }

if sys.platform == "win32":
    warnings.filterwarnings("ignore", category=DeprecationWarning)
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

try:
    from styling import inject_custom_css
    from ui_modules import render_hero_banner, render_marquee_header
    from registry import initialize_system_states, HON_TITLE
    import panels
except Exception as init_err:
    st.error(f"🛑 CRITICAL INITIALIZATION ERROR: Failed to import sub-modules.")
    st.exception(init_err)
    st.stop()

# --- 2. PAGE CONFIG & STYLING ---
st.set_page_config(
    page_title="LSOEP - Honourable Victor Abang, PhD Portal",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- 2.1 UNIVERSAL MULTI-FORMAT DOCUMENT READER ---
def read_word_document(file_path):
    if not os.path.exists(file_path):
        return None
        
    is_odt = False
    try:
        with open(file_path, 'rb') as f:
            header = f.read(4)
            if b'PK' in header: 
                f.seek(30)
                if b'opendocument.text' in f.read(100):
                    is_odt = True
    except Exception:
        pass

    # Route 1: ODT Parser (LibreOffice native or masked ODT format)
    if is_odt or file_path.endswith('.odt'):
        try:
            from odf import text, teletype
            from odf.opendocument import load
            odt_doc = load(file_path)
            paragraphs = odt_doc.getElementsByType(text.P)
            full_text = []
            for p in paragraphs:
                val = teletype.extractText(p).strip()
                if val:
                    full_text.append(val)
            return "\n\n".join(full_text)
        except Exception as odt_err:
            st.error(f"LibreOffice extraction engine failed: {str(odt_err)}")
            return None

    # Route 2: Standard Microsoft Word DOCX parser
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
            
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))
                    
        return "\n\n".join(full_text)
    except Exception as e:
        # Fallback in case of custom masked structures
        try:
            from odf import text, teletype
            from odf.opendocument import load
            odt_doc = load(file_path)
            paragraphs = odt_doc.getElementsByType(text.P)
            return "\n\n".join([teletype.extractText(p).strip() for p in paragraphs if teletype.extractText(p).strip()])
        except Exception:
            st.error("⚠️ FORMAT ERROR: Your file structure does not match a valid Office XML (.docx) or LibreOffice (.odt) structure.")
            return None

# --- 3. RUNTIME EXECUTION WRAPPER ---
try:
    inject_custom_css()
    initialize_system_states()

    # --- 4. NAVIGATION MATRICES (Grouped Logically for Fast Navigation) ---
    PORTAL_CATEGORIES = {
        "📜 Legislative & Plenary": [
            "🏛️ LEGISLATIVE FOOTPRINTS",
            "🚀 LEGISLATIVE PROGRESS TRACKER",
            "🏛️ LIVE PLENARY UPDATES",
            "🏛️ BEYOND RHETORICS PROJECT EXECUTION"
        ],
        "🎓 Empowerment & Grants": [
            "🛠️ SKILL VOCATION POOL",
            "🎓 STUDENT SCHOLARSHIP/GRANT",
            "📦 PALLIATIVE ENROLLMENT",
            "💡 CV & ARTISAN VAULT"
        ],
        "🗣️ Feedback & Security": [
            "🚨 COMMUNITY URGENT NEED",
            "🗣️ SPEAK TO HON. VICTOR ABANG DIRECTLY",
            "🛡️ LOCAL LEADERSHIP VOUCHING"
        ]
    }
    
    # Flattened list for routing branches
    NAVIGATION_OPTIONS = [item for sublist in PORTAL_CATEGORIES.values() for item in sublist]
    
    ADMIN_OPTIONS = {
        "CONTROL_ROOM": "🔑 EXECUTIVE CONTROL ROOM",
        "STRATEGIC_COMMITTEES": "🛡️ STRATEGIC COMMITTEES (MODULE 13)",
        "AGENT_HUB": "🗳️ POLLING UNIT AGENT HUB",
        "COLLATION_HUB": "🛡️ WARD COLLATION OFFICER HUB",
    }

    # --- 5. UNIVERSAL SIDEBAR ARCHITECTURE WITH INTERACTIVE BLUE FB ICON LOGO ---
    st.sidebar.markdown(
        f"""
<a href="https://www.facebook.com/profile.php?id=100076989890731" target="_blank" style="text-decoration:none;">
<div style="display:flex; align-items:center; justify-content:center; gap:10px; padding:12px; background-color:#061A33; border:2px solid #1877F2; border-radius:8px; color:#1877F2; font-weight:bold; transition: 0.3s; box-shadow: 0 4px 10px rgba(24,119,242,0.15);">
<svg width="22" height="22" viewBox="0 0 24 24" fill="#1877F2" xmlns="http://www.w3.org/2000/svg">
<path d="M24 12.073c0-6.627-5.373-12-12-12s-12 5.373-12 12c0 5.99 4.388 10.954 10.125 11.854v-8.385H7.078v-3.47h3.047V9.43c0-3.007 1.792-4.669 4.533-4.669 1.312 0 2.686.235 2.686.235v2.953H15.83c-1.491 0-1.956.925-1.956 1.874v2.25h3.328l-.532 3.47h-2.796v8.385C19.612 23.027 24 18.062 24 12.073z"/>
</svg>
<span style="color:#FFFFFF; font-size: 0.95rem; letter-spacing:0.3px;">Hon. Victor Abang Official Page</span>
</div>
</a>
""",
        unsafe_allow_html=True,
    )
    st.sidebar.markdown("<hr style='margin:15px 0; border-color:#0B3C5D;'>", unsafe_allow_html=True)
    st.sidebar.markdown("<h3 style='color:#D4AF37; font-size:1.1rem; margin-bottom:10px; text-transform:uppercase; letter-spacing:0.5px;'>Admin Checkpoints</h3>", unsafe_allow_html=True)

    # Admin buttons
    if st.sidebar.button(ADMIN_OPTIONS["CONTROL_ROOM"], use_container_width=True, key="nav_btn_admin"):
        st.session_state.current_route = ADMIN_OPTIONS["CONTROL_ROOM"]
        st.rerun()
    if st.sidebar.button(ADMIN_OPTIONS["STRATEGIC_COMMITTEES"], use_container_width=True, key="nav_btn_committee"):
        st.session_state.current_route = ADMIN_OPTIONS["STRATEGIC_COMMITTEES"]
        st.rerun()
    if st.sidebar.button(ADMIN_OPTIONS["AGENT_HUB"], use_container_width=True, key="nav_btn_agent"):
        st.session_state.current_route = ADMIN_OPTIONS["AGENT_HUB"]
        st.rerun()
    if st.sidebar.button(ADMIN_OPTIONS["COLLATION_HUB"], use_container_width=True, key="nav_btn_collation"):
        st.session_state.current_route = ADMIN_OPTIONS["COLLATION_HUB"]
        st.rerun()

    # --- SIDEBAR CONTROL ROOM NAVIGATION MODULES ---
    if st.session_state.current_route == ADMIN_OPTIONS["CONTROL_ROOM"] and st.session_state.admin_authenticated:
        st.sidebar.markdown("<hr style='margin:15px 0; border-color:#0B3C5D;'>", unsafe_allow_html=True)
        st.sidebar.markdown("<h3 style='color:#D4AF37; font-size:1.1rem; margin-bottom:12px; text-transform:uppercase; letter-spacing:0.5px;'>Command Hub Tabs</h3>", unsafe_allow_html=True)
        
        admin_modules_list = [
            "📊 Master Registry Matrix",
            "📢 Plenary Broadcast Terminal",
            "🗣️ Citizen Feedback",
            "📢 Admin Announcement Control",
            "⚖️ Database Audit Diagnostics",
            "🛡️ RADAR Deduplication Interceptor",
            "🎓 Scholar Talent Matrix",
            "💎 Vantedge Influencer Proportions",
            "🗳️ Live Election Analytical Sync",
            "📝 Ground Truth Form EC8A Data",
            "📂 Bulk Data Sync Stream",
            "📜 Executive Waiver Ledger",
            "🚀 Legislative Progress Tracker",
            "📅 Long-Term Momentum Monitoring",
            "📋 Strategic Committee Compliance Logs",
        ]
        
        for mod in admin_modules_list:
            is_active = st.session_state.admin_module_view == mod
            bg_color = "#0B3C5D" if is_active else "#020B14"
            border_style = "2px solid #D4AF37" if is_active else "1px solid #0B3C5D"
            
            st.sidebar.markdown(
                f'''<div style="padding:2px; margin-bottom:6px; background-color:{bg_color}; border:{border_style}; border-radius:6px; text-align:center;">''', 
                unsafe_allow_html=True
            )
            if st.sidebar.button(mod, key=f"side_mod_btn_{mod}", use_container_width=True):
                st.session_state.admin_module_view = mod
                st.rerun()
            st.sidebar.markdown("</div>", unsafe_allow_html=True)

    st.sidebar.caption("Architecture Build: v75.1 | Premium Tiles Active")

    # --- 6. GLOBAL ROUTER & LAYOUT ENGINE ---
    selected_route = st.session_state.current_route
    st.caption("🌐 Platform Stream Active")

    if selected_route == "HOME":
        render_hero_banner()
        
        # 🏢 THE DEDICATED LAYOUT OVERLAY SPACER
        st.markdown("<div style='margin-top: 30px;'></div>", unsafe_allow_html=True)
        
        render_marquee_header()
        
        # --- 🏛️ PREMIUM PULSE BOX HEADER CONTAINER ---
        st.markdown(
            """
<style>
@keyframes institutional-pulse {
0% { box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.4), 0 0 0 0 rgba(214, 175, 55, 0.2); transform: scale(1); }
50% { box-shadow: 0 12px 40px 0 rgba(0, 0, 0, 0.5), 0 0 20px 6px rgba(214, 175, 55, 0.4); transform: scale(1.015); }
100% { box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.4), 0 0 0 0 rgba(214, 175, 55, 0.2); transform: scale(1); }
}
.premium-pulse-box {
background: linear-gradient(135deg, rgba(6, 26, 51, 0.75) 0%, rgba(3, 20, 36, 0.9) 100%);
padding: 24px 35px;
border-radius: 12px;
border: 2px solid rgba(214, 175, 55, 0.4);
text-align: center;
margin-top: 25px !important;
margin-bottom: 25px !important;
animation: institutional-pulse 3s infinite ease-in-out;
}
.premium-pulse-text {
background: linear-gradient(135deg, #FFF6D6 0%, #D4AF37 50%, #AA7C11 100%);
-webkit-background-clip: text;
-webkit-text-fill-color: transparent;
font-family: 'Inter', system-ui, sans-serif;
font-weight: 900 !important;
font-size: 2.1rem !important;
letter-spacing: 3px;
text-transform: uppercase;
margin: 0 !important;
padding: 0 !important;
}
</style>
<div class="premium-pulse-box">
<h2 class="premium-pulse-text">🏛️ CONSTITUENCY PORTAL CONTROL DECK</h2>
</div>
""", 
            unsafe_allow_html=True
        )
        
        # --- 🚀 THE ELITE CATEGORIZED TAB SWITCHER ENGINE ---
        st.markdown(
            """
<style>
/* Target all buttons inside our portal navigation section */
div[data-testid="stColumn"] button {
    background: linear-gradient(135deg, #0b2240 0%, #061526 100%) !important;
    color: #f8fafc !important;
    border: 2px solid rgba(212, 175, 55, 0.5) !important;
    border-radius: 8px !important;
    padding: 14px 20px !important;
    font-size: 16px !important;
    font-weight: 700 !important;
    letter-spacing: 0.5px !important;
    text-transform: uppercase !important;
    transition: all 0.3s ease-in-out !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.25) !important;
    min-height: 58px !important;
    margin-bottom: 12px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}

/* Hover effects for a tactile, glowing feel with a unique text color transformation */
div[data-testid="stColumn"] button:hover {
    border-color: #D4AF37 !important;
    background: linear-gradient(135deg, #0e2a4f 0%, #0a213a 100%) !important;
    box-shadow: 0 0 15px rgba(212, 175, 55, 0.4) !important;
    transform: translateY(-2px) !important;
    color: #FFE896 !important;
    text-shadow: 0px 0px 8px rgba(255, 232, 150, 0.6) !important;
}

/* Click/Active state animation */
div[data-testid="stColumn"] button:active {
    transform: translateY(1px) !important;
    box-shadow: 0 2px 5px rgba(0, 0, 0, 0.5) !important;
}

/* Column headers styling */
.category-header {
    font-family: 'Inter', sans-serif;
    color: #D4AF37;
    font-size: 1.3rem;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin-bottom: 5px;
}
</style>
""",
            unsafe_allow_html=True
        )

        cols = st.columns(3)
        for col_idx, (category_name, options) in enumerate(PORTAL_CATEGORIES.items()):
            with cols[col_idx]:
                st.markdown(f'<p class="category-header">{category_name}</p>', unsafe_allow_html=True)
                st.markdown("<hr style='border-color: rgba(212, 175, 55, 0.3); margin: 5px 0 20px 0;'>", unsafe_allow_html=True)
                
                for opt in options:
                    clean_name = opt
                    if st.button(clean_name, key=f"portal_nav_{opt}", use_container_width=True):
                        st.session_state.current_route = opt
                        st.rerun()
        
    else:
        render_marquee_header()
        if st.button("↩️ Return to Main Gateway", use_container_width=True, key="nav_btn_return"):
            st.session_state.current_route = "HOME"
            st.session_state.admin_authenticated = False
            st.session_state.agent_authenticated = False
            st.rerun()
        st.markdown("<hr class='nav-divider'>", unsafe_allow_html=True)

        # --- Panel Routing Branches ---
        if selected_route == "🏛️ LEGISLATIVE FOOTPRINTS":
            panels.render_sponsored_bills_panel()
        elif selected_route == "🚀 LEGISLATIVE PROGRESS TRACKER":
            panels.render_progress_tracker()
        elif selected_route == "🛠️ SKILL VOCATION POOL":
            panels.render_skill_form()
        elif selected_route == "🎓 STUDENT SCHOLARSHIP/GRANT":
            panels.render_scholarship_form()
            
        # --- REDIRECTS & INTEGRATED PALLIATIVE MODULE ---
        elif selected_route == "📦 PALLIATIVE ENROLLMENT":
            panels.render_palliative_form(focus_on_vouching=False)
            
        elif selected_route == "🛡️ LOCAL LEADERSHIP VOUCHING":
            panels.render_palliative_form(focus_on_vouching=True)
            
        # --- BEYOND RHETORICS INTERFACE (WITH CROSS-PLATFORM FALLBACK) ---
        elif selected_route == "🏛️ BEYOND RHETORICS PROJECT EXECUTION":
            st.markdown("## 🦅 BEYOND RHETORICS: PROJECT VERIFICATION HUB")
            st.caption("Cross-examining performance metrics with verifiable ground-truth evidence.")
            
            # --- CROSS-PLATFORM PATH RESOLVER ---
            if "USERPROFILE" in os.environ:
                DESKTOP_PATH = os.path.join(os.environ["USERPROFILE"], "Desktop")
            elif "HOME" in os.environ:
                DESKTOP_PATH = os.path.join(os.environ["HOME"], "Desktop")
                if not os.path.exists(DESKTOP_PATH):
                    DESKTOP_PATH = os.environ["HOME"]
            else:
                DESKTOP_PATH = os.getcwd()

            DOC_FILENAME = "Hon_Victor_Abang_Projects.docx"
            FULL_DOC_PATH = os.path.join(DESKTOP_PATH, DOC_FILENAME)

            if os.path.exists(FULL_DOC_PATH):
                with open(FULL_DOC_PATH, "rb") as f:
                    doc_bytes = f.read()

                col1, col2 = st.columns([3, 1])
                with col1:
                    st.success(f"Successfully connected to active dataset: **{DOC_FILENAME}**")
                with col2:
                    st.download_button(
                        label="📥 Download Original DOCX",
                        data=doc_bytes,
                        file_name=DOC_FILENAME,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                with st.spinner("Extracting verification metrics..."):
                    document_content = read_word_document(FULL_DOC_PATH)

                if document_content:
                    st.info("📊 **Verifiable Performance Metrics Index**")
                    st.markdown(
                        f"""
<div style="background-color: #0f172a; border: 1px solid #1e293b; padding: 25px; border-radius: 8px; max-height: 600px; overflow-y: auto; color: #f8fafc; font-family: 'Inter', -apple-system, sans-serif; font-size: 18px; line-height: 1.6; letter-spacing: 0.2px;">
{document_content.replace('\n', '<br>')}
</div>
""", 
                        unsafe_allow_html=True
                    )
            else:
                st.warning(
                    f"Verified project documentation file was not found on your desktop at: `{FULL_DOC_PATH}`. "
                    "Please save your LibreOffice / Word document on your Desktop with the exact filename to activate the verified stream."
                )
                
        elif selected_route == "🗣️ SPEAK TO HON. VICTOR ABANG DIRECTLY":
            panels.render_speak_directly_panel()
        elif selected_route == "🏛️ LIVE PLENARY UPDATES":
            panels.render_constituent_plenary_updates()

        # --- Admin Channels ---
        elif selected_route == ADMIN_OPTIONS["CONTROL_ROOM"]:
            if st.session_state.admin_authenticated:
                panels.main_dashboard(conn=None)
            else:
                st.markdown("### 🔑 Executive Command System Authorization")
                admin_key = st.text_input("Enter Command Hub Key:", type="password", key="admin_key_input")
                if st.button("Authorize Access", key="admin_auth_button", use_container_width=True):
                    if admin_key == "victor2027":
                        st.session_state.admin_authenticated = True
                        st.rerun()
                    elif admin_key:
                        st.error("🛑 SYSTEM ACCESS REJECTED")
                        
        elif selected_route == ADMIN_OPTIONS["STRATEGIC_COMMITTEES"]:
            panels.strategic_committees_panel()
            
        elif selected_route == ADMIN_OPTIONS["AGENT_HUB"]:
            if st.session_state.get("agent_authenticated", False):
                panels.agent_panel()
            else:
                st.markdown("### 🦅 Polling Unit Agent Security Checkpoint")
                agent_key = st.text_input("Enter Polling Unit Agent Authority Key:", type="password", key="gate_agent_key_input")
                if st.button("Authorize Agent Hub Access", key="agent_auth_execute_btn", use_container_width=True):
                    if agent_key == "victor2027":
                        st.session_state.agent_authenticated = True
                        st.success("Authorization successful! Opening Agent Command Node...")
                        st.rerun()
                    elif agent_key:
                        st.error("🛑 ACCESS REJECTED: Invalid Polling Agent Authority Signature.")

        elif selected_route == ADMIN_OPTIONS["COLLATION_HUB"]:
            st.markdown("### 🛡️ Ward Collation Command Security Checkpoint")
            collation_key = st.text_input("Enter Collation Officer Key:", type="password", key="gate_collation_key")
            if st.button("Authorize Hub Access", key="collation_auth_btn", use_container_width=True):
                if collation_key == "victor2027":
                    panels.ward_collation_officer_panel()
                elif collation_key:
                    st.error("🛑 ACCESS REJECTED: Invalid Collation Authority Signature.")

except Exception as runtime_err:
    st.error("🛑 RUNTIME ENGINE CRASH INTERCEPTED")
    st.exception(runtime_err)